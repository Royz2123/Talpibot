# -*- coding: iso-8859-8 -*-

from shutil import copyfile
import datetime

from docx import Document
from docx.shared import Inches

import template

class HantarDocument(object):
    HANTAR_BASE_NAME = 'outputs\\' + 'hantar_report - '
    DEFAULT_TEMPLATE = "templates\\template.docx"

    def __init__(self, filename=None):
        # todays date
        self._date = datetime.datetime.today()
            
        # output filename
        self._filename = filename
        if filename is None:
            self._filename = "%s%s.docx" % (
                HantarDocument.HANTAR_BASE_NAME,
                self._date.strftime('%Y-%m-%d')
            )
                        
        # copy the template file
        copyfile(HantarDocument.DEFAULT_TEMPLATE, self._filename)        
            
        # template we are copying
        self._object = template.TemplateDocument(self._filename)
       
    def fill_document(self):
        pars = self._object.get_paragraphs()
        
        print(self._object.print_paragraphs())
        
        pars[0].add_run(self._date.strftime('%Y-%m-%d'))        
        
        pars[2].add_run(self._date.strftime('%Y-%m-%d'))        
        
        
        self._object.save()
    
    EDITS = []
    
    
    
    
    
    
   
   
   

        
    

    