from docx import Document
from docx.shared import Inches

class TemplateDocument(object):
    def __init__(self, filename):
        self._filename = filename
        self._template_document = Document(filename)

    def get_paragraphs(self):
        return self._template_document.paragraphs

    def print_paragraphs(self):
        for index, p in enumerate(self._template_document.paragraphs):
            print("%d:\t:%s" % (index, p.text))
            
    def save(self):
        self._template_document.save(self._filename)